from typing import List, Tuple, Optional, TypedDict
import html as html_module
import time
import random
import unicodedata
import urllib.parse

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import RGBColor

EDITORIAL_INCLUDED_NOTE = "in existing editorially included list"
EDITORIAL_EXCLUDED_NOTE = "in existing editorially excluded list"

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

OED_HOME = "https://www.oed.com/"

# -----------------------------
# Helpers
# -----------------------------
def parse_word_list(text: str) -> List[str]:
    words = []
    for line in text.splitlines():
        word = line.strip()
        if word:
            words.append(unicodedata.normalize("NFC", word).upper())
    return words


HYPHEN_CHARS = "-‐‑‒–—―"


def contains_hyphen(text: str) -> bool:
    return any(ch in HYPHEN_CHARS for ch in text)


def contains_accent(text: str) -> bool:
    normalized = unicodedata.normalize("NFD", text)
    if any(unicodedata.category(ch) == "Mn" for ch in normalized):
        return True
    return any(ch.isalpha() and ord(ch) > 127 for ch in text)


def get_word_form_rejection(word: str) -> Optional[str]:
    if contains_hyphen(word):
        return "Contains hyphen"
    if contains_accent(word):
        return "Contains accent"
    return None


def detect_center_letter(words: List[str]) -> str:
    if not words:
        return ""

    letter_sets = [set(word.upper()) for word in words]
    common_letters = set.intersection(*letter_sets)

    if len(common_letters) == 1:
        return next(iter(common_letters))
    if len(common_letters) > 1:
        return ", ".join(sorted(common_letters))
    return ""


class PuzzleLevels(TypedDict):
    wa: int
    wahoo: int
    wahoowa: int
    wahoo_wow: int
    average: int


def calculate_levels(total_words: int, words: Optional[List[str]] = None) -> PuzzleLevels:
    if total_words <= 0:
        return {
            "wa": 0,
            "wahoo": 0,
            "wahoowa": 0,
            "wahoo_wow": 0,
            "average": 0,
        }

    wahoo_wow = total_words

    # Wa: low bar that most players can reach (~12% of puzzle, minimum 5 words).
    wa = max(5, round(total_words * 0.12))

    # Wahoo: encouraging next step (~30% of puzzle, at least 5 words above Wa).
    wahoo = max(wa + 5, round(total_words * 0.30))
    wahoo = min(wahoo, max(wa + 1, wahoo_wow - 1))
    wa = min(wa, max(1, wahoo - 1))

    # Average: midpoint between Wa and Wahoo.
    average = round((wa + wahoo) / 2)
    average = max(wa + 1, min(average, wahoo - 1))

    # Wahoowa: estimate from shorter, familiar-length words in the final list.
    if words:
        easy_word_count = sum(1 for word in words if 4 <= len(word) <= 6)
        wahoowa = easy_word_count if easy_word_count > 0 else round(total_words * 0.45)
    else:
        wahoowa = round(total_words * 0.45)

    wahoowa = max(wahoo + 1, min(wahoowa, wahoo_wow))

    return {
        "wa": wa,
        "wahoo": wahoo,
        "wahoowa": wahoowa,
        "wahoo_wow": wahoo_wow,
        "average": average,
    }


def apply_wahoowa_override(levels: PuzzleLevels, wahoowa_override: Optional[int]) -> PuzzleLevels:
    if wahoowa_override is None or wahoowa_override <= 0:
        return levels

    updated = dict(levels)
    updated["wahoowa"] = max(
        levels["wahoo"] + 1,
        min(wahoowa_override, levels["wahoo_wow"]),
    )
    return updated  # type: ignore[return-value]


def apply_editorial_filters(
    common: List[Tuple[str, str, str]],
    rare: List[Tuple[str, str, str]],
    editorial_included: List[str],
    editorial_excluded: List[str],
) -> Tuple[List[Tuple[str, str, str]], List[Tuple[str, str, str]]]:
    common_dict = {word: (word, link, reason) for word, link, reason in common}
    rare_dict = {word: (word, link, reason) for word, link, reason in rare}

    # Force-include: only override OED rejections for words on the editorial list.
    for word in editorial_included:
        if word in rare_dict:
            _, link, _ = rare_dict.pop(word)
            common_dict[word] = (word, link, EDITORIAL_INCLUDED_NOTE)

    # Force-exclude: only override OED approvals for words on the editorial list.
    for word in editorial_excluded:
        if word in common_dict:
            _, link, _ = common_dict.pop(word)
            rare_dict[word] = (word, link, EDITORIAL_EXCLUDED_NOTE)

    common = sorted(common_dict.values(), key=lambda x: x[0])
    rare = sorted(rare_dict.values(), key=lambda x: x[0])
    return common, rare


# -----------------------------
# OED functions
# -----------------------------
def oed_link(word: str) -> str:
    return f"{OED_HOME}search/dictionary/?scope=Entries&q={urllib.parse.quote(word)}"


def _classify_oed_result(word: str, result, url: str) -> Tuple[str, str, str]:
    link = f'<a href="{url}">{word}</a>'

    try:
        title = result.find_element(By.CLASS_NAME, "hw").text.strip()
    except Exception:
        title = ""

    if contains_hyphen(title):
        return "rare", link, "Hyphenated form"
    if contains_accent(title):
        return "rare", link, "Accented form"

    try:
        freq_div = result.find_element(By.CLASS_NAME, "frequencyIndicator")
        usage = int(freq_div.get_attribute("aria-valuenow"))
    except Exception:
        usage = None

    try:
        snippet = result.find_element(By.CLASS_NAME, "snippet").text.lower()
    except Exception:
        snippet = ""

    try:
        ps_text = result.find_element(By.CLASS_NAME, "ps").text.lower()
    except Exception:
        ps_text = ""

    is_variant = "variant of" in ps_text
    is_slang = "slang" in snippet

    is_proper_noun = (
        bool(title)
        and title[0].isupper()
        and any(c.islower() for c in title[1:])
    )

    if is_variant:
        return "rare", link, "Variant form"
    if is_proper_noun:
        return "rare", link, "Proper noun"
    if is_slang:
        return "rare", link, "Slang"
    if usage is None:
        return "rare", link, "Missing frequency data"
    if usage <= 2:
        return "rare", link, f"Low frequency ({usage})"
    return "common", link, ""


def classify_words(words: List[str]):
    common, rare = [], []

    options = Options()
    options.add_argument("--disable-gpu")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")

    driver = webdriver.Chrome(options=options)
    wait = WebDriverWait(driver, 6)

    for word in words:
        print(f"Processing {word}...")

        rejection = get_word_form_rejection(word)
        if rejection:
            rare.append((word, f'<a href="{oed_link(word)}">{word}</a>', rejection))
            continue

        time.sleep(random.uniform(4, 7))
        url = oed_link(word)

        try:
            driver.get(url)

            try:
                result = wait.until(
                    EC.visibility_of_element_located((By.CLASS_NAME, "resultsSetItem"))
                )
            except Exception:
                rare.append((word, f'<a href="{url}">{word}</a>', "No results"))
                continue

            bucket, link, reason = _classify_oed_result(word, result, url)
            if bucket == "common":
                common.append((word, link, reason))
            else:
                rare.append((word, link, reason))

        except Exception:
            rare.append((word, f'<a href="{oed_link(word)}">{word}</a>', "Error fetching"))

    driver.quit()

    common.sort(key=lambda x: x[0])
    rare.sort(key=lambda x: x[0])

    return common, rare


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def split_removed_words(rare: List[Tuple[str, str, str]]):
    removed_words = []
    editorially_removed = []
    for item in rare:
        if item[2] == EDITORIAL_EXCLUDED_NOTE:
            editorially_removed.append(item)
        else:
            removed_words.append(item)
    return removed_words, editorially_removed


def add_bold_heading(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    return paragraph


def add_word_entry(doc: Document, word: str, reason: str = "", numbered: bool = True):
    style = "List Number" if numbered else None
    paragraph = doc.add_paragraph(style=style)
    add_hyperlink(paragraph, oed_link(word), word)
    if reason:
        paragraph.add_run(f" – {reason}")


def create_docx(
    common,
    rare,
    filename="oed_classification.docx",
    creator: str = "",
    link: str = "",
    levels: Optional[PuzzleLevels] = None,
    puzzle_title: str = "",
):
    doc = Document()

    word_count = len(common)
    center_letter = detect_center_letter([word for word, _, _ in common])
    removed_words, editorially_removed = split_removed_words(rare)
    if levels is None:
        levels = calculate_levels(word_count, [word for word, _, _ in common])

    if puzzle_title:
        doc.add_heading(puzzle_title, 0)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    metadata = [
        ("Creator", creator),
        ("Center Letter", center_letter),
        ("Link", link),
    ]
    for row_index, (label, value) in enumerate(metadata):
        table.rows[row_index].cells[0].text = label
        table.rows[row_index].cells[1].text = value

    doc.add_paragraph()

    add_bold_heading(doc, "Words")
    for word, _, reason in common:
        add_word_entry(doc, word, reason)

    doc.add_paragraph()

    add_bold_heading(doc, "Removed Words")
    for word, _, reason in removed_words:
        add_word_entry(doc, word, reason)

    doc.add_paragraph()

    editorial_header = doc.add_paragraph()
    editorial_run = editorial_header.add_run("Editorially Removed")
    editorial_run.font.color.rgb = RGBColor(0, 0, 255)
    editorial_run.font.underline = True
    editorial_header.add_run(" Words ")
    flag_run = editorial_header.add_run("FLAG FOR COPY")
    flag_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    for word, _, reason in editorially_removed:
        add_word_entry(doc, word, reason)
    if not editorially_removed:
        doc.add_paragraph("-")

    doc.add_paragraph()

    add_bold_heading(doc, "Levels")
    doc.add_paragraph(f"Wa = {levels['wa']}")
    doc.add_paragraph(f"Wahoo = {levels['wahoo']}")
    doc.add_paragraph(f"Wahoowa = {levels['wahoowa']}")
    doc.add_paragraph(f"WahooWOW = {levels['wahoo_wow']}")
    doc.add_paragraph(f"Average = {levels['average']}")

    doc.save(filename)
    return filename


def write_html_file(title: str, items):
    html = f"""<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"><title>{title}</title></head>
<body>
<h2>{title}</h2>
<ul>
"""
    html += "\n".join(
        f"<li>{link} – {reason}</li>" for (_, link, reason) in items
    )
    html += "\n</ul></body></html>"
    return html


def _escape(text: str) -> str:
    return html_module.escape(text or "")


TABLE_CELL_STYLE = "border: 1px solid #000000; padding: 4px 8px;"


def _table_cell(content: str) -> str:
    return f'<td style="{TABLE_CELL_STYLE}">{content}</td>'


def _table_row(label: str, value: str) -> str:
    return f"<tr>{_table_cell(_escape(label))}{_table_cell(value)}</tr>"


def _word_list_item_html(word: str, reason: str = "") -> str:
    line = f'<a href="{_escape(oed_link(word))}">{_escape(word)}</a>'
    if reason:
        line += f" – {_escape(reason)}"
    return f"<li>{line}</li>"


def _word_list_html(items) -> str:
    if not items:
        return "<ol><li>-</li></ol>"
    return "<ol>" + "".join(_word_list_item_html(word, reason) for word, _, reason in items) + "</ol>"


def build_output_html(
    common,
    rare,
    creator: str = "",
    link: str = "",
    levels: Optional[PuzzleLevels] = None,
    puzzle_title: str = "",
    center_letter: str = "",
) -> str:
    removed_words, editorially_removed = split_removed_words(rare)
    if levels is None:
        levels = calculate_levels(len(common), [word for word, _, _ in common])

    parts = [
        "<!DOCTYPE html>",
        "<html>",
        "<head>",
        '<meta charset="UTF-8">',
        "<style>",
        "body { font-family: Arial, sans-serif; }",
        "h1 { margin-bottom: 16px; }",
        "p { margin: 4px 0; }",
        "ol { margin-top: 4px; }",
        "</style>",
        "</head>",
        "<body>",
    ]

    if puzzle_title:
        parts.append(f"<h1>{_escape(puzzle_title)}</h1>")

    parts.append(
        '<table border="1" cellspacing="0" cellpadding="4" '
        'style="border-collapse: collapse; width: 100%; margin-bottom: 16px;">'
    )
    metadata = [
        ("Creator", _escape(creator)),
        ("Center Letter", _escape(center_letter)),
        ("Link", f'<a href="{_escape(link)}">{_escape(link)}</a>' if link else ""),
    ]
    for label, value in metadata:
        parts.append(_table_row(label, value))
    parts.append("</table>")

    parts.append("<p><strong>Words</strong></p>")
    parts.append(_word_list_html(common))

    parts.append("<p><strong>Removed Words</strong></p>")
    parts.append(_word_list_html(removed_words))

    parts.append("<p>")
    parts.append('<span style="color: blue; text-decoration: underline;">Editorially Removed</span>')
    parts.append(' Words <span style="background: yellow;">FLAG FOR COPY</span></p>')

    parts.append(_word_list_html(editorially_removed))

    parts.append("<p><strong>Levels</strong></p>")
    level_labels = [
        ("wa", "Wa"),
        ("wahoo", "Wahoo"),
        ("wahoowa", "Wahoowa"),
        ("wahoo_wow", "WahooWOW"),
        ("average", "Average"),
    ]
    for key, label in level_labels:
        parts.append(f"<p>{label} = {levels[key]}</p>")

    parts.append("</body></html>")
    return "".join(parts)
